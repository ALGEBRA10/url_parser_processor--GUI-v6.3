#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
🌐 URL处理器 v6.3 - 精确提取漏洞URL（优化版：支持纯IP/IP:端口 + Markdown URL，勾选时按域名去重）
"""
import ctypes
ctypes.windll.user32.ShowWindow(ctypes.windll.kernel32.GetConsoleWindow(), 0)
import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext
import re
import os
import threading
import requests
from urllib.parse import urlparse
try:
    from docx import Document
except ImportError:
    messagebox.showerror("错误", "请安装依赖: pip install python-docx requests")
    exit()

# 公共后缀
SPECIAL_TLDS = {
    'com.cn', 'net.cn', 'org.cn', 'gov.cn', 'ac.cn',
    'co.uk', 'me.uk', 'org.uk', 'com.au', 'net.au',
}

# Markdown 风格链接正则：[文本](https://example.com/xxx)
MD_LINK_RE = re.compile(r'\[([^\]]+)\]\((https?://[^\s)]+)\)')

def extract_url_from_markdown(line: str) -> str:
    """
    从一行中提取真正的 URL:
    - 优先匹配 [text](url) 结构，返回 url
    - 如果没有 markdown 结构，则原样返回
    """
    line = line.strip()
    m = MD_LINK_RE.search(line)
    if m:
        # m.group(2) 为真正的链接
        return m.group(2).strip()
    return line


class URLProcessorApp:
    def __init__(self, root):
        self.root = root
        self.root.title("🌐 URL处理器 v6.3 - 支持纯IP/IP:端口")
        self.root.geometry("1200x800")
        self.root.resizable(True, True)
        self.records = []
        # 默认不去重，只在勾选时按域名去重
        self.dedup_enabled = False
        self.setup_ui()

    def setup_ui(self):
        main_frame = ttk.Frame(self.root)
        main_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        ctrl_frame = ttk.LabelFrame(main_frame, text="控制面板", padding=5)
        ctrl_frame.pack(fill=tk.X, pady=5)

        ttk.Button(ctrl_frame, text="📂 选择目录(.docx)", command=self.load_directory).pack(side=tk.LEFT, padx=5)
        
        self.dedup_var = tk.BooleanVar(value=False)
        ttk.Checkbutton(
            ctrl_frame,
            text="按域名去重",
            variable=self.dedup_var,
            command=self.on_toggle_dedup
        ).pack(side=tk.LEFT, padx=10)

        ttk.Button(ctrl_frame, text="➡ 输入框解析", command=self.load_from_textbox).pack(side=tk.LEFT, padx=5)
        ttk.Button(ctrl_frame, text="🔍 HTTP存活检测", command=self.check_http_alive).pack(side=tk.RIGHT, padx=5)
        ttk.Button(ctrl_frame, text="🗑️ 清空", command=self.clear_all).pack(side=tk.RIGHT, padx=5)

        self.path_var = tk.StringVar(value="就绪")
        ttk.Label(ctrl_frame, textvariable=self.path_var, foreground="gray").pack(side=tk.LEFT, padx=10)

        input_frame = ttk.LabelFrame(main_frame, text="手动输入URL（支持纯IP/IP:端口，每行一个）", padding=5)
        input_frame.pack(fill=tk.X, pady=(0, 5))
        self.input_text = scrolledtext.ScrolledText(input_frame, height=4, wrap=tk.NONE)
        self.input_text.pack(fill=tk.BOTH, expand=True)

        paned = ttk.Panedwindow(main_frame, orient=tk.VERTICAL)
        paned.pack(fill=tk.BOTH, expand=True, pady=5)

        # 表格
        result_frame = ttk.LabelFrame(paned, text="解析结果", padding=5)
        paned.add(result_frame, weight=3)

        columns = ("原始隐患URL", "子域名", "域名", "WEB站点URL", "HTTP状态")
        self.tree = ttk.Treeview(result_frame, columns=columns, show='headings', height=18)
        
        widths = {"原始隐患URL": 380, "子域名": 160, "域名": 160, "WEB站点URL": 260, "HTTP状态": 80}
        for col in columns:
            self.tree.heading(col, text=col)
            self.tree.column(col, width=widths[col], anchor='w')

        v_scroll = ttk.Scrollbar(result_frame, orient=tk.VERTICAL, command=self.tree.yview)
        h_scroll = ttk.Scrollbar(result_frame, orient=tk.HORIZONTAL, command=self.tree.xview)
        self.tree.configure(yscrollcommand=v_scroll.set, xscrollcommand=h_scroll.set)
        self.tree.grid(row=0, column=0, sticky='nsew')
        v_scroll.grid(row=0, column=1, sticky='ns')
        h_scroll.grid(row=1, column=0, sticky='ew')
        result_frame.rowconfigure(0, weight=1)
        result_frame.columnconfigure(0, weight=1)

        # 四块输出
        text_frame = ttk.LabelFrame(paned, text="分块结果（可复制）", padding=5)
        paned.add(text_frame, weight=2)

        text_frame.columnconfigure(0, weight=1)
        text_frame.columnconfigure(1, weight=1)
        text_frame.rowconfigure(0, weight=1)
        text_frame.rowconfigure(1, weight=1)

        def make_block(parent, row, col, title):
            frame = ttk.Frame(parent)
            frame.grid(row=row, column=col, sticky="nsew", padx=5, pady=5)
            ttk.Label(frame, text=title).pack(anchor="w")
            text = scrolledtext.ScrolledText(frame, height=7, wrap=tk.NONE)
            text.pack(fill=tk.BOTH, expand=True, pady=2)
            text.config(state=tk.DISABLED)
            ttk.Button(frame, text="复制全部", command=lambda: self.copy_text(text, title)).pack(anchor="e", pady=2)
            return text

        self.text_domain = make_block(text_frame, 0, 0, "域名")
        self.text_subdomain = make_block(text_frame, 0, 1, "子域名")
        self.text_weburl = make_block(text_frame, 1, 0, "WEB站点URL")
        self.text_original = make_block(text_frame, 1, 1, "原始隐患URL")

        self.progress_var = tk.DoubleVar(value=0)
        self.progress_bar = ttk.Progressbar(main_frame, orient="horizontal", mode="determinate", variable=self.progress_var)
        self.progress_bar.pack(fill=tk.X, pady=(0, 5))

        self.status_var = tk.StringVar(value="就绪")
        ttk.Label(main_frame, textvariable=self.status_var, relief=tk.SUNKEN, anchor=tk.W).pack(fill=tk.X, pady=(5, 0))

    def copy_text(self, text_widget, title):
        self.root.clipboard_clear()
        self.root.clipboard_append(text_widget.get("1.0", tk.END).strip())
        messagebox.showinfo("复制成功", f"已复制 {title} 全部内容")

    def on_toggle_dedup(self):
        # 勾选 -> 启用按域名去重
        self.dedup_enabled = self.dedup_var.get()

    def _init_progress(self, maximum):
        self.progress_bar["maximum"] = max(1, maximum)
        self.progress_var.set(0)
        self.progress_bar.update()

    def _update_progress(self, current, total=None):
        if total:
            self.progress_bar["maximum"] = max(1, total)
        self.progress_var.set(current)
        self.progress_bar.update()

    def _finish_progress(self):
        self.progress_var.set(0)

    def load_directory(self):
        dir_path = filedialog.askdirectory(title="选择包含 .docx 的目录")
        if not dir_path:
            return

        self.path_var.set(f"目录: {os.path.basename(dir_path)}")
        self.status_var.set("正在扫描目录...")
        self.root.update()

        files = [
            (os.path.join(dir_path, name), os.path.getmtime(os.path.join(dir_path, name)))
            for name in os.listdir(dir_path)
            if name.lower().endswith(".docx")
        ]

        if not files:
            messagebox.showinfo("提示", "目录中没有 .docx 文件")
            return

        files.sort(key=lambda x: x[1])
        all_urls = []
        total = len(files)
        self._init_progress(total)

        for idx, (file_path, _) in enumerate(files, 1):
            try:
                urls = self.read_docx_vuln_urls(file_path)
                all_urls.extend(urls)
                self.status_var.set(
                    f"扫描进度: {idx}/{total} | {os.path.basename(file_path)} | 累计 {len(all_urls)} 个URL"
                )
            except Exception as e:
                self.status_var.set(f"读取失败: {str(e)}")
            self._update_progress(idx, total)
            self.root.update()

        self._finish_progress()
        self.status_var.set(f"完成扫描，共提取 {len(all_urls)} 个URL")
        if all_urls:
            self.process_urls(all_urls)
        else:
            messagebox.showinfo("提示", "未找到漏洞URL，尝试手动输入测试")

    def read_docx_vuln_urls(self, file_path):
        """精确提取：只取'漏洞URL'等标签后面的第一个有效URL（支持 Markdown 链接）"""
        urls = []
        try:
            doc = Document(file_path)
            
            vuln_patterns = [
                r'漏洞URL[:：]?\s*([a-zA-Z0-9:/.\-\[\]\(\)]+?)(?=\s|$)',
                r'漏洞地址[:：]?\s*([a-zA-Z0-9:/.\-\[\]\(\)]+?)(?=\s|$)',
                r'目标[:：]?\s*([a-zA-Z0-9:/.\-\[\]\(\)]+?)(?=\s|$)',
                r'URL[:：]?\s*([a-zA-Z0-9:/.\-\[\]\(\)]+?)(?=\s|$)'
            ]
            
            all_text = ""
            for para in doc.paragraphs:
                all_text += para.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        all_text += cell.text + "\n"
            
            found_urls = set()
            for pattern in vuln_patterns:
                match = re.search(pattern, all_text, re.IGNORECASE)
                if match:
                    raw = match.group(1).strip()
                    url = extract_url_from_markdown(raw)
                    if self.is_valid_url(url) and url not in found_urls:
                        urls.append(url)
                        found_urls.add(url)
            
            return urls
        except:
            return []

    def is_valid_ip(self, host):
        """验证IPv4"""
        try:
            parts = host.split('.')
            return len(parts) == 4 and all(0 <= int(p) <= 255 for p in parts)
        except:
            return False

    def is_valid_url(self, text):
        """验证URL（支持协议、纯IP、IP:端口、域名:端口）"""
        text = text.strip().lower()
        if any(text.startswith(p) for p in ['http://', 'https://', 'mysql://', 'redis://']):
            return True
        if '://' in text:
            return True
        if re.match(r'^\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3}(:\d+)?$', text):
            return True
        if re.match(r'^[a-z0-9][a-z0-9-]*(\.[a-z0-9][a-z0-9-]*)+\.?(:\d+)?$', text):
            return True
        return False

    def extract_domain(self, host):
        """提取域名（纯IP保持原样）"""
        host = host.strip('.').lower()
        if self.is_valid_ip(host):
            return host
        
        if ':' in host and self.is_valid_ip(host.split(':')[0]):
            return host.split(':')[0]
            
        parts = host.split('.')
        if len(parts) < 2:
            return host
        
        for i in range(len(parts) - 1, 1, -1):
            suffix = '.'.join(parts[i:])
            if suffix in SPECIAL_TLDS:
                return '.'.join(parts[i-1:])
        return '.'.join(parts[-2:])

    def parse_url(self, raw_url):
        """核心解析函数 - 支持所有情况 + Markdown 预处理"""
        original_input = raw_url.strip()
        if not original_input:
            return None

        # 先处理 Markdown 包装
        url_clean = extract_url_from_markdown(original_input)
        original = url_clean.strip()
        if not original:
            return None
        
        # 1. 有协议的特殊服务（如mysql://, redis://）
        if '://' in original and not original.lower().startswith(('http://', 'https://')):
            parts = original.split('://')
            if len(parts) > 1:
                host_part = parts[1].split('/')[0].split('@')[-1].split(':')[0]
                return {
                    "原始隐患URL": original_input, 
                    "子域名": host_part, 
                    "域名": self.extract_domain(host_part), 
                    "WEB站点URL": host_part
                }
        
        # 2. 纯IP 或 IP:端口
        ip_match = re.match(r'^(\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3})(:(\d+))?$', original)
        if ip_match:
            ip = ip_match.group(1)
            port = f":{ip_match.group(3)}" if ip_match.group(3) else ""
            web_url = f"http://{ip}{port}"  # 自动补协议用于HTTP检测
            return {
                "原始隐患URL": original_input,
                "子域名": f"{ip}{port}",
                "域名": ip,
                "WEB站点URL": web_url
            }
        
        # 3. HTTP/HTTPS 或域名
        if not original.lower().startswith(('http://', 'https://')):
            original = 'http://' + original
        
        try:
            parsed = urlparse(original)
            host = parsed.hostname or ""
            if not host:
                return None
            port = f":{parsed.port}" if parsed.port else ""
            subdomain = f"{host}{port}"
            return {
                "原始隐患URL": original_input,
                "子域名": subdomain,
                "域名": self.extract_domain(host),
                "WEB站点URL": f"{parsed.scheme}://{host}{port}"
            }
        except:
            return None

    def load_from_textbox(self):
        """手动输入：严格一行一个URL，自动识别 Markdown 链接"""
        raw = self.input_text.get("1.0", tk.END).strip()
        if not raw:
            messagebox.showinfo("提示", "请先输入URL")
            return
        
        urls = []
        for line in raw.splitlines():
            line = line.strip()
            if not line:
                continue
            pure = extract_url_from_markdown(line)
            if pure and self.is_valid_url(pure):
                urls.append(pure)
        
        if not urls:
            messagebox.showinfo("提示", "未检测到有效URL，请检查格式")
            return
            
        self.path_var.set(f"输入框：{len(urls)} 条")
        self.process_urls(urls)

    def process_urls(self, urls):
        # 清空界面
        for item in self.tree.get_children():
            self.tree.delete(item)
        self.records.clear()
        
        for txt in [self.text_subdomain, self.text_domain, self.text_weburl, self.text_original]:
            txt.config(state=tk.NORMAL)
            txt.delete("1.0", tk.END)
        
        # 按「域名」去重（仅在勾选时）
        seen_domains = set()
        count = 0
        
        for url in urls:
            record = self.parse_url(url)
            if not record:
                continue

            domain = record["域名"].lower()

            # 只有勾选复选框时才去重
            if self.dedup_enabled and domain in seen_domains:
                continue
            if self.dedup_enabled:
                seen_domains.add(domain)
            
            item_id = self.tree.insert("", tk.END, values=(
                record["原始隐患URL"], record["子域名"], record["域名"], 
                record["WEB站点URL"], ""
            ))
            record["tree_item_id"] = item_id
            record["http_status"] = ""
            self.records.append(record)
            
            self.text_original.insert(tk.END, record["原始隐患URL"] + "\n")
            self.text_subdomain.insert(tk.END, record["子域名"] + "\n")
            self.text_domain.insert(tk.END, record["域名"] + "\n")
            self.text_weburl.insert(tk.END, record["WEB站点URL"] + "\n")
            count += 1
        
        for txt in [self.text_subdomain, self.text_domain, self.text_weburl, self.text_original]:
            txt.config(state=tk.DISABLED)
        
        self.root.title(f"🌐 URL处理器 v6.3 | {count} 条 (输入{len(urls)}条)")
        self.status_var.set(f"解析完成：输入{len(urls)}条，展示{count}条")

    def clear_all(self):
        for item in self.tree.get_children():
            self.tree.delete(item)
        self.records.clear()
        for txt in [self.text_subdomain, self.text_domain, self.text_weburl, self.text_original]:
            txt.config(state=tk.NORMAL)
            txt.delete("1.0", tk.END)
            txt.config(state=tk.DISABLED)
        self.input_text.delete("1.0", tk.END)
        self.path_var.set("就绪")
        self.status_var.set("已清空")
        self._finish_progress()

    def _update_http_status(self, item_id, status_text):
        if not item_id:
            return
        values = list(self.tree.item(item_id, "values"))
        values[4] = status_text
        self.tree.item(item_id, values=values)

    def check_http_alive(self):
        if not self.records:
            messagebox.showinfo("提示", "请先加载URL")
            return
        threading.Thread(target=self._check_http_alive_worker, daemon=True).start()

    def _check_http_alive_worker(self):
        total = len(self.records)
        self.status_var.set(f"HTTP检测中... {total}条")
        self._init_progress(total)
        
        for idx, rec in enumerate(self.records, 1):
            url = rec.get("WEB站点URL", "")
            status = "非HTTP"
            if url and url.startswith(('http://', 'https://')):
                try:
                    r = requests.get(url, timeout=5, allow_redirects=True)
                    status = "存活" if 200 <= r.status_code < 300 else f"异常({r.status_code})"
                except:
                    status = "失败"
            
            rec["http_status"] = status
            self.root.after(0, self._update_http_status, rec["tree_item_id"], status)
            self._update_progress(idx, total)
        
        self.status_var.set("HTTP检测完成")
        self._finish_progress()


def main():
    root = tk.Tk()
    app = URLProcessorApp(root)
    root.mainloop()


if __name__ == "__main__":
    main()
